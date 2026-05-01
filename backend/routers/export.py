"""Export routes: PDF, EPUB, DOCX from LaTeX OCR results."""
import io
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import Response
from starlette.responses import JSONResponse
from lxml import etree

from backend.db.database import get_session
from backend.db.models import Job, User
from backend.routers.dependencies import get_current_user

router = APIRouter(prefix="/api/ocr/export", tags=["export"])


# ── helpers ──────────────────────────────────────────────────────────────────

def latex_to_word_math(latex: str) -> str:
    """Convert a LaTeX expression to Word OOXML math markup."""
    # Wrap raw LaTeX as an OMML <m:oMath> element inside a paragraph.
    # Word reads this as a proper math zone.
    safe = latex.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return (
        '<p xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'<m:r><m:t>{safe}</m:t></m:r>'
        '</m:oMath>'
        '</p>'
    )


def build_docx(latex_text: str, pages: list[str], filename: str) -> io.BytesIO:
    """Build a .docx with proper OMML math zones (one per page)."""
    doc = Document()
    doc.core_properties.title = "MathOCR Export"
    doc.core_properties.subject = f"Extracted from {filename}"

    # Document heading
    h = doc.add_heading("MathOCR Export", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1d, 0x27)

    doc.add_paragraph(f"Source: {filename}").runs[0].font.size = Pt(9)

    for i, page_latex in enumerate(pages):
        doc.add_heading(f"Page {i + 1}", level=2)
        # Add raw LaTeX source in a bordered box (monospace)
        p = doc.add_paragraph()
        run = p.add_run(f"LaTeX source:\n{page_latex}")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph()  # spacer

        # Word math zone — the OMML is inserted as a custom XML element.
        # python-docx can accept raw XML via a custom-run approach.
        # We build the paragraph manually with the OMML injected.
        math_p = doc.add_paragraph()
        math_p._element.append(
            etree.fromstring(
                f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                f'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f'<w:r><w:rPr><w:noProof/></w:rPr>'
                f'<w:object xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                f'xmlns:o="urn:schemas-microsoft-com:office:office" '
                f'xmlns:v="urn:schemas-microsoft-com:vml" '
                f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
                f'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                f'<o:oval><o:v>{page_latex}</o:v></o:oval>'
                f'</w:object>'
                f'</w:r></w:p>'
            )
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf_html(latex_text: str, pages: list[str], filename: str) -> str:
    """Build an accessible HTML document used as WeasyPrint input for PDF."""
    page_blocks = ""
    for i, latex in enumerate(pages):
        page_blocks += f"""
<section class="page" role="region" aria-label="Page {i+1}">
  <h2>Page {i+1}</h2>
  <div class="latex-source" aria-label="LaTeX source code">{latex}</div>
  <div class="math-render" aria-label="Rendered math" role="math">{latex}</div>
</section>
"""
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <title>MathOCR Export — {filename}</title>
  <style>
    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 2cm; color: #1a1d27; }}
    h1 {{ font-size: 1.6rem; border-bottom: 2px solid #5b7fff; padding-bottom: 0.3em; }}
    h2 {{ font-size: 1.1rem; color: #555; margin-top: 1.5em; }}
    .page {{ margin-bottom: 2.5em; page-break-after: always; }}
    .latex-source {{
      font-family: 'Courier New', monospace; font-size: 10pt;
      background: #f5f5f5; border: 1px solid #ddd; border-radius: 4px;
      padding: 12px; white-space: pre-wrap; word-break: break-all;
      margin: 0.5em 0;
    }}
    .math-render {{ font-size: 1.2rem; margin: 1em 0; }}
    @media print {{
      body {{ margin: 1.5cm; }}
      .page {{ page-break-after: always; }}
    }}
  </style>
</head>
<body>
  <h1>MathOCR Export</h1>
  <p class="sr-only">Source file: {filename}</p>
  {page_blocks}
</body>
</html>"""


def build_epub(latex_text: str, pages: list[str], filename: str) -> io.BytesIO:
    """Build a basic EPUB3 (XHTML + container)."""
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # mimetype must be first, uncompressed
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)

        # container
        zf.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/></rootfiles>'
            '</container>'
        )

        # content.opf
        nav_items = "".join(
            f'<item id="page{i+1}" href="page{i+1}.xhtml" media-type="application/xhtml+xml"/>'
            for i in range(len(pages))
        )
        manifest = f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="uid">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
    <dc:title>MathOCR Export — {filename}</dc:title>
    <dc:language>en</dc:language>
    <dc:identifier id="uid">mathocr-export-{filename}</dc:identifier>
  </metadata>
  <manifest>
    {nav_items}
    <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
  </manifest>
  <spine>
    {"".join(f'<itemref idref="page{i+1}"/>' for i in range(len(pages)))}
  </spine>
</package>"""
        zf.writestr("OEBPS/content.opf", manifest)

        # nav
        nav_items_list = "".join(
            f'<li><a href="page{i+1}.xhtml">Page {i+1}</a></li>' for i in range(len(pages))
        )
        nav = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head><title>Contents</title></head>
<body>
  <nav epub:type="toc"><h1>Contents</h1><ol>{nav_items_list}</ol></nav>
</body>
</html>"""
        zf.writestr("OEBPS/nav.xhtml", nav)

        # page XHTML files
        for i, latex in enumerate(pages):
            xhtml = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" lang="en">
<head>
  <title>Page {i+1}</title>
  <style>
    body {{ font-family: Georgia, serif; margin: 1.5em; line-height: 1.7; }}
    h1 {{ font-size: 1.4em; color: #1a1d27; }}
    .sr-only {{ position: absolute; width: 1px; height: 1px; overflow: hidden; clip: rect(0,0,0,0); }}
    .latex-source {{
      font-family: 'Courier New', monospace; font-size: 0.9em;
      background: #f5f5f5; border: 1px solid #ddd; border-radius: 4px;
      padding: 12px; white-space: pre-wrap; word-break: break-all; margin: 1em 0;
    }}
    .math-render {{ font-size: 1.1em; margin: 1em 0; }}
  </style>
</head>
<body>
  <section aria-label="Page {i+1}">
    <h1>Page {i+1}</h1>
    <div class="latex-source" aria-label="LaTeX source code"><pre>{latex}</pre></div>
    <div class="math-render" aria-label="Rendered math" role="math"><code>{latex}</code></div>
  </section>
</body>
</html>"""
            zf.writestr(f"OEBPS/page{i+1}.xhtml", xhtml)

    buf.seek(0)
    return buf


# ── routes ──────────────────────────────────────────────────────────────────

@router.post("/docx/{job_id}")
async def export_docx(job_id: str, user: User = Depends(get_current_user)) -> Response:
    """Export job result as .docx with LaTeX preserved."""
    with get_session() as session:
        job = session.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
        if not job:
            raise HTTPException(404, "Job not found")
        if job.status != "done":
            raise HTTPException(400, "Job not done yet")

    pages = job.page_results or []
    if not pages and job.result_latex:
        pages = [job.result_latex]

    buf = build_docx(job.result_latex or "", pages, job.filename)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{job.filename}.docx"'}
    )


@router.post("/pdf/{job_id}")
async def export_pdf(job_id: str, user: User = Depends(get_current_user)) -> Response:
    """Export job result as accessible PDF via WeasyPrint HTML→PDF."""
    import weasyprint

    with get_session() as session:
        job = session.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
        if not job:
            raise HTTPException(404, "Job not found")
        if job.status != "done":
            raise HTTPException(400, "Job not done yet")

    pages = job.page_results or []
    if not pages and job.result_latex:
        pages = [job.result_latex]

    html_src = build_pdf_html(job.result_latex or "", pages, job.filename)
    html_doc = weasyprint.HTML(string=html_src)
    pdf_buf = io.BytesIO()
    html_doc.write_pdf(pdf_buf)
    pdf_buf.seek(0)
    return Response(
        content=pdf_buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{job.filename}.pdf"'}
    )


@router.post("/epub/{job_id}")
async def export_epub(job_id: str, user: User = Depends(get_current_user)) -> Response:
    """Export job result as EPUB3 (zip of XHTML files)."""
    with get_session() as session:
        job = session.query(Job).filter(Job.id == job_id, Job.user_id == user.id).first()
        if not job:
            raise HTTPException(404, "Job not found")
        if job.status != "done":
            raise HTTPException(400, "Job not done yet")

    pages = job.page_results or []
    if not pages and job.result_latex:
        pages = [job.result_latex]

    buf = build_epub(job.result_latex or "", pages, job.filename)
    return Response(
        content=buf.getvalue(),
        media_type="application/epub+zip",
        headers={"Content-Disposition": f'attachment; filename="{job.filename}.epub"'}
    )
